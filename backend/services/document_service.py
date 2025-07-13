import os
import uuid
from typing import List, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import Config

class DocumentService:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""]
        )
    
    def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """保存上传的文件并返回文件路径"""
        file_id = str(uuid.uuid4())
        file_extension = Path(filename).suffix
        saved_filename = f"{file_id}{file_extension}"
        file_path = os.path.join(Config.UPLOAD_FOLDER, saved_filename)
        
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        return file_path
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """从PDF文件提取文本"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"PDF文件解析失败: {str(e)}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """从DOCX文件提取文本"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            raise Exception(f"DOCX文件解析失败: {str(e)}")
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """从TXT文件提取文本"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    text = file.read()
            except Exception as e:
                raise Exception(f"TXT文件解析失败: {str(e)}")
        except Exception as e:
            raise Exception(f"TXT文件解析失败: {str(e)}")
        return text
    
    def extract_text(self, file_path: str) -> str:
        """根据文件类型提取文本"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise Exception(f"不支持的文件类型: {file_extension}")
    
    def split_text_to_chunks(self, text: str, filename: str) -> List[LangchainDocument]:
        """将文本分割成块"""
        if not text.strip():
            raise Exception("文档内容为空")
        
        # 使用LangChain的文本分割器
        chunks = self.text_splitter.split_text(text)
        
        # 创建LangChain文档对象
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": filename,
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                }
            )
            documents.append(doc)
        
        return documents
    
    def process_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """处理文档：保存、提取文本、分块"""
        try:
            # 检查文件类型
            file_extension = Path(filename).suffix.lower()
            if file_extension not in ['.pdf', '.txt', '.docx']:
                raise Exception(f"不支持的文件类型: {file_extension}")
            
            # 保存文件
            file_path = self.save_uploaded_file(file_content, filename)
            
            # 提取文本
            text = self.extract_text(file_path)
            
            # 检查文本长度
            if len(text.strip()) < 10:
                raise Exception("文档内容太少，请上传包含更多内容的文档")
            
            # 分割文本
            documents = self.split_text_to_chunks(text, filename)
            
            # 清理临时文件
            os.remove(file_path)
            
            return {
                "success": True,
                "filename": filename,
                "text_length": len(text),
                "chunk_count": len(documents),
                "documents": documents,
                "preview": text[:500] + "..." if len(text) > 500 else text
            }
            
        except Exception as e:
            # 清理可能创建的临时文件
            if 'file_path' in locals() and os.path.exists(file_path):
                os.remove(file_path)
            
            return {
                "success": False,
                "error": str(e)
            }
    
    def get_document_summary(self, documents: List[LangchainDocument]) -> Dict[str, Any]:
        """获取文档摘要信息"""
        total_chars = sum(len(doc.page_content) for doc in documents)
        
        return {
            "total_chunks": len(documents),
            "total_characters": total_chars,
            "average_chunk_size": total_chars // len(documents) if documents else 0,
            "sources": list(set(doc.metadata.get("source", "") for doc in documents))
        } 